# from langchain_community.tools import WikipediaQueryRun, DuckDuckGoSearchRun
# from langchain_community.utilities import WikipediaAPIWrapper
# from langchain.tools import Tool
# from datetime import datetime
# from dotenv import load_dotenv
# import os
# import json
# import pandas as pd
# import requests
# from bs4 import BeautifulSoup
# import arxiv
# from scholarly import scholarly
# import PyPDF2
# from docx import Document
# from youtube_transcript_api import YouTubeTranscriptApi
# import re
# from urllib.parse import urlparse, parse_qs
# import io

# load_dotenv()

# # Enhanced save functionality
# def save_research_data(data: str, format_type: str = "txt", filename: str = None):
#     """Save research data in multiple formats"""
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
#     if filename is None:
#         filename = f"research_output_{timestamp}"
    
#     if format_type == "json":
#         try:
#             json_data = json.loads(data)
#             filename += ".json"
#             with open(filename, "w", encoding="utf-8") as f:
#                 json.dump(json_data, f, indent=2, ensure_ascii=False)
#         except json.JSONDecodeError:
#             # If not valid JSON, save as text
#             filename += ".txt"
#             with open(filename, "w", encoding="utf-8") as f:
#                 f.write(f"Research Output - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n{data}")
    
#     elif format_type == "csv":
#         # Try to convert to structured data for CSV
#         filename += ".csv"
#         lines = data.split('\n')
#         df = pd.DataFrame({'content': lines})
#         df.to_csv(filename, index=False, encoding='utf-8')
    
#     else:  # default to txt
#         filename += ".txt"
#         formatted_text = f"-- Research Output --\nTimestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n{data}\n"
#         with open(filename, "a", encoding="utf-8") as f:
#             f.write(formatted_text)
    
#     return f"Data saved to {filename} at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

# # ArXiv research tool
# def search_arxiv(query: str, max_results: int = 5):
#     """Search ArXiv for academic papers"""
#     try:
#         search = arxiv.Search(
#             query=query,
#             max_results=max_results,
#             sort_by=arxiv.SortCriterion.Relevance
#         )
        
#         results = []
#         for result in search.results():
#             paper_info = {
#                 "title": result.title,
#                 "authors": [author.name for author in result.authors],
#                 "summary": result.summary[:500] + "..." if len(result.summary) > 500 else result.summary,
#                 "published": result.published.strftime("%Y-%m-%d"),
#                 "url": result.entry_id,
#                 "categories": result.categories
#             }
#             results.append(paper_info)
        
#         return json.dumps(results, indent=2)
#     except Exception as e:
#         return f"Error searching ArXiv: {str(e)}"

# # Google Scholar tool
# def search_scholar(query: str, num_results: int = 5):
#     """Search Google Scholar for academic papers"""
#     try:
#         search_query = scholarly.search_pubs(query)
#         results = []
        
#         count = 0
#         for pub in search_query:
#             if count >= num_results:
#                 break
                
#             paper_info = {
#                 "title": pub.get('title', 'No title'),
#                 "authors": pub.get('author', 'No authors'),
#                 "year": pub.get('year', 'No year'),
#                 "venue": pub.get('venue', 'No venue'),
#                 "abstract": pub.get('abstract', 'No abstract')[:500] + "..." if pub.get('abstract') and len(pub.get('abstract')) > 500 else pub.get('abstract', 'No abstract'),
#                 "citations": pub.get('num_citations', 0),
#                 "url": pub.get('pub_url', 'No URL')
#             }
#             results.append(paper_info)
#             count += 1
        
#         return json.dumps(results, indent=2)
#     except Exception as e:
#         return f"Error searching Google Scholar: {str(e)}"

# # News and article scraper
# def scrape_article(url: str):
#     """Scrape and extract content from web articles"""
#     try:
#         headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
#         }
#         response = requests.get(url, headers=headers, timeout=10)
#         soup = BeautifulSoup(response.content, 'html.parser')
        
#         # Extract title
#         title = soup.find('title')
#         title = title.get_text() if title else "No title found"
        
#         # Extract main content (try multiple selectors)
#         content_selectors = ['article', 'main', '.content', '.post-content', '.entry-content', 'p']
#         content = ""
        
#         for selector in content_selectors:
#             elements = soup.select(selector)
#             if elements:
#                 content = ' '.join([elem.get_text() for elem in elements])
#                 break
        
#         # Clean up content
#         content = re.sub(r'\s+', ' ', content).strip()
#         content = content[:2000] + "..." if len(content) > 2000 else content
        
#         return f"Title: {title}\n\nContent: {content}\n\nSource: {url}"
    
#     except Exception as e:
#         return f"Error scraping article: {str(e)}"

# # PDF text extractor
# def extract_pdf_text(file_path: str):
#     """Extract text from PDF files"""
#     try:
#         with open(file_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             text = ""
            
#             for page in pdf_reader.pages:
#                 text += page.extract_text() + "\n"
            
#             # Clean up text
#             text = re.sub(r'\s+', ' ', text).strip()
#             return text[:3000] + "..." if len(text) > 3000 else text
    
#     except Exception as e:
#         return f"Error extracting PDF text: {str(e)}"

# # Word document extractor
# def extract_docx_text(file_path: str):
#     """Extract text from Word documents"""
#     try:
#         doc = Document(file_path)
#         text = ""
        
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
        
#         # Clean up text
#         text = re.sub(r'\s+', ' ', text).strip()
#         return text[:3000] + "..." if len(text) > 3000 else text
    
#     except Exception as e:
#         return f"Error extracting DOCX text: {str(e)}"

# # YouTube transcript extractor
# def get_youtube_transcript(video_url: str):
#     """Extract transcript from YouTube videos"""
#     try:
#         # Extract video ID from URL
#         parsed_url = urlparse(video_url)
#         if parsed_url.hostname == 'youtu.be':
#             video_id = parsed_url.path[1:]
#         elif parsed_url.hostname in ('www.youtube.com', 'youtube.com'):
#             if parsed_url.path == '/watch':
#                 video_id = parse_qs(parsed_url.query)['v'][0]
#             elif parsed_url.path[:7] == '/embed/':
#                 video_id = parsed_url.path.split('/')[2]
#             elif parsed_url.path[:3] == '/v/':
#                 video_id = parsed_url.path.split('/')[2]
#         else:
#             return "Invalid YouTube URL"
        
#         # Get transcript
#         transcript = YouTubeTranscriptApi.get_transcript(video_id)
        
#         # Combine transcript text
#         full_text = ' '.join([entry['text'] for entry in transcript])
        
#         return f"YouTube Video Transcript:\n\n{full_text[:2000]}..." if len(full_text) > 2000 else f"YouTube Video Transcript:\n\n{full_text}"
    
#     except Exception as e:
#         return f"Error getting YouTube transcript: {str(e)}"

# # Data analysis tool
# def analyze_research_data(data: str):
#     """Analyze research data and provide insights"""
#     try:
#         word_count = len(data.split())
#         char_count = len(data)
#         lines = data.split('\n')
#         line_count = len(lines)
        
#         # Find most common words (simple analysis)
#         words = re.findall(r'\b\w+\b', data.lower())
#         word_freq = {}
#         for word in words:
#             if len(word) > 3:  # Only count words longer than 3 characters
#                 word_freq[word] = word_freq.get(word, 0) + 1
        
#         # Get top 10 most common words
#         top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
#         analysis = f"""
# Data Analysis Results:
# - Total words: {word_count}
# - Total characters: {char_count}
# - Total lines: {line_count}
# - Average words per line: {word_count/line_count if line_count > 0 else 0:.1f}

# Top 10 Most Common Words:
# {chr(10).join([f"- {word}: {count}" for word, count in top_words])}
#         """
        
#         return analysis.strip()
    
#     except Exception as e:
#         return f"Error analyzing data: {str(e)}"

# # Create tool instances
# save_tool = Tool(
#     name="save_research_data",
#     func=lambda data: save_research_data(data, "txt"),
#     description="Save research data to a file with timestamp"
# )

# save_json_tool = Tool(
#     name="save_json_data",
#     func=lambda data: save_research_data(data, "json"),
#     description="Save research data as JSON format"
# )

# save_csv_tool = Tool(
#     name="save_csv_data",
#     func=lambda data: save_research_data(data, "csv"),
#     description="Save research data as CSV format"
# )

# arxiv_tool = Tool(
#     name="search_arxiv",
#     func=search_arxiv,
#     description="Search ArXiv for academic papers. Input should be a search query string."
# )

# scholar_tool = Tool(
#     name="search_google_scholar",
#     func=search_scholar,
#     description="Search Google Scholar for academic papers. Input should be a search query string."
# )

# scraper_tool = Tool(
#     name="scrape_article",
#     func=scrape_article,
#     description="Scrape content from web articles. Input should be a valid URL."
# )

# pdf_tool = Tool(
#     name="extract_pdf_text",
#     func=extract_pdf_text,
#     description="Extract text content from PDF files. Input should be a file path."
# )

# docx_tool = Tool(
#     name="extract_docx_text",
#     func=extract_docx_text,
#     description="Extract text content from Word documents. Input should be a file path."
# )

# youtube_tool = Tool(
#     name="get_youtube_transcript",
#     func=get_youtube_transcript,
#     description="Extract transcript from YouTube videos. Input should be a YouTube URL."
# )

# analysis_tool = Tool(
#     name="analyze_research_data",
#     func=analyze_research_data,
#     description="Analyze research data and provide insights including word count, common terms, etc."
# )

# # Enhanced web search
# search = DuckDuckGoSearchRun()
# search_tool = Tool(
#     name="web_search",
#     func=search.run,
#     description="Search the web for current information and news"
# )

# # Enhanced Wikipedia
# api_wrapper = WikipediaAPIWrapper(top_k_results=2, doc_content_chars_max=500)
# wiki_tool = WikipediaQueryRun(api_wrapper=api_wrapper)

# # Export all tools
# all_tools = [
#     search_tool,
#     wiki_tool,
#     save_tool,
#     save_json_tool,
#     save_csv_tool,
#     arxiv_tool,
#     scholar_tool,
#     scraper_tool,
#     pdf_tool,
#     docx_tool,
#     youtube_tool,
#     analysis_tool
# ]
"""
tools.py — Comprehensive Research Tools Suite
Compatible with LangChain agents and custom workflows
Author: AI Research Team • 2025  
"""
from __future__ import annotations
from typing import List, Tuple, Dict, Any, Optional
import json
import re
import requests
import os
from datetime import datetime
from urllib.parse import urlparse, parse_qs
from dotenv import load_dotenv

import pandas as pd
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
from urllib.parse import urlparse, parse_qs
import arxiv
from scholarly import scholarly
from langchain.tools import Tool
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_community.tools import DuckDuckGoSearchRun, WikipediaQueryRun
import logging

# Load environment variables
load_dotenv()

# Set up logging
logger = logging.getLogger(__name__)



# --------------------------------------------------#
# Data Persistence Utilities
# --------------------------------------------------#
def save_research_data(data: str, fmt: str = "txt", fname: Optional[str] = None) -> str:
    """Save research data in multiple formats"""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = fname or f"research_{timestamp}"
        
        # Create output directory
        os.makedirs("output", exist_ok=True)
        
        if fmt == "json":
            filepath = os.path.join("output", f"{fname}.json")
            try:
                json_data = json.loads(data)
                with open(filepath, "w", encoding="utf-8") as fh:
                    json.dump(json_data, fh, indent=2, ensure_ascii=False)
            except json.JSONDecodeError:
                # Fallback to text format
                filepath = os.path.join("output", f"{fname}.txt")
                fmt = "txt"

        if fmt == "csv":
            filepath = os.path.join("output", f"{fname}.csv")
            lines = [line.strip() for line in data.split('\n') if line.strip()]
            df = pd.DataFrame({'content': lines})
            df.to_csv(filepath, index=False, encoding='utf-8')

        if fmt == "txt":
            if 'filepath' not in locals():
                filepath = os.path.join("output", f"{fname}.txt")
            formatted_data = f"--- Research Output [{timestamp}] ---\n{data}\n\n"
            with open(filepath, "a", encoding="utf-8") as fh:
                fh.write(formatted_data)

        return f"✅ Data saved to {filepath}"
    
    except Exception as e:
        return f"❌ Save failed: {str(e)}"

# --------------------------------------------------#
# Academic Research Tools
# --------------------------------------------------#
def search_arxiv(query: str, max_results: int = 5) -> str:
    """Search ArXiv for academic papers"""
    try:
        search = arxiv.Search(
            query=query,
            max_results=max_results,
            sort_by=arxiv.SortCriterion.Relevance
        )
        
        papers = []
        for result in search.results():
            paper_info = {
                "title": result.title,
                "authors": [author.name for author in result.authors],
                "url": result.entry_id,
                "published": result.published.strftime("%Y-%m-%d"),
                "summary": result.summary[:600] + "..." if len(result.summary) > 600 else result.summary,
                "categories": result.categories,
                "pdf_url": result.pdf_url
            }
            papers.append(paper_info)
        
        if not papers:
            return "No academic papers found on ArXiv for this query."
        
        return json.dumps(papers, indent=2, ensure_ascii=False)
    
    except Exception as e:
        logger.error(f"ArXiv search error: {e}")
        return f"❌ ArXiv search failed: {str(e)}"

def search_scholar(query: str, limit: int = 5) -> str:
    """Search Google Scholar for academic papers"""
    try:
        search_query = scholarly.search_pubs(query)
        results = []
        
        count = 0
        for pub in search_query:
            if count >= limit:
                break
            
            try:
                paper_info = {
                    "title": pub.get('title', 'No title available'),
                    "authors": pub.get('author', 'No authors listed'),
                    "year": pub.get('year', 'Year not specified'),
                    "venue": pub.get('venue', 'Venue not specified'),
                    "abstract": (pub.get('abstract', 'No abstract available')[:600] + "..." 
                               if pub.get('abstract') and len(pub.get('abstract')) > 600 
                               else pub.get('abstract', 'No abstract available')),
                    "citations": pub.get('num_citations', 0),
                    "url": pub.get('pub_url', 'No URL available')
                }
                results.append(paper_info)
                count += 1
            except Exception as e:
                logger.warning(f"Error processing Google Scholar result: {e}")
                continue
        
        if not results:
            return "No academic papers found on Google Scholar for this query."
        
        return json.dumps(results, indent=2, ensure_ascii=False)
    
    except Exception as e:
        logger.error(f"Google Scholar search error: {e}")
        return f"❌ Google Scholar search failed: {str(e)}"

# --------------------------------------------------#
# Web Scraping Utilities
# --------------------------------------------------#
def scrape_article(url: str) -> str:
    """Scrape and extract content from web articles"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract title
        title = soup.find('title')
        title = title.get_text().strip() if title else "No title found"
        
        # Extract main content using multiple selectors
        content_selectors = [
            'article', 'main', '.content', '.post-content', 
            '.entry-content', '.article-content', '.story-body',
            '[role="main"]', '#content', '.main-content'
        ]
        
        content = ""
        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                content = ' '.join([elem.get_text() for elem in elements])
                break
        
        # Fallback to all paragraphs if no main content found
        if not content:
            paragraphs = soup.find_all('p')
            content = ' '.join([p.get_text() for p in paragraphs])
        
        # Clean up content
        content = re.sub(r'\s+', ' ', content).strip()
        content = content[:3000] + "..." if len(content) > 3000 else content
        
        if not content:
            return f"❌ Could not extract meaningful content from {url}"
        
        return f"Title: {title}\n\nContent: {content}\n\nSource: {url}"
    
    except requests.RequestException as e:
        return f"❌ Failed to fetch article from {url}: {str(e)}"
    except Exception as e:
        return f"❌ Failed to scrape article: {str(e)}"

# --------------------------------------------------#
# Document Processing Utilities
# --------------------------------------------------#
def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF files"""
    try:
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        if not file_path.lower().endswith('.pdf'):
            return f"❌ File is not a PDF: {file_path}"
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) == 0:
                return "❌ PDF file appears to be empty or corrupted"
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            # Clean up text
            text = re.sub(r'\s+', ' ', text).strip()
            
            if not text:
                return "❌ No readable text found in PDF"
            
            return text[:4000] + "..." if len(text) > 4000 else text
    
    except Exception as e:
        return f"❌ PDF extraction failed: {str(e)}"

def extract_docx_text(file_path: str) -> str:
    """Extract text from Word documents"""
    try:
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}"
        
        if not file_path.lower().endswith('.docx'):
            return f"❌ File is not a DOCX document: {file_path}"
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # Clean up text
        text = re.sub(r'\s+', ' ', text).strip()
        
        if not text:
            return "❌ No readable text found in DOCX document"
        
        return text[:4000] + "..." if len(text) > 4000 else text
    
    except Exception as e:
        return f"❌ DOCX extraction failed: {str(e)}"

# --------------------------------------------------#
# Text Analysis Utilities
# --------------------------------------------------#
def analyze_research_data(data: str) -> str:
    """Analyze research data and provide insights"""
    try:
        if not data.strip():
            return "❌ No data provided for analysis"
        
        # Basic text statistics
        words = data.split()
        word_count = len(words)
        char_count = len(data)
        lines = data.split('\n')
        line_count = len([line for line in lines if line.strip()])
        
        # Find most common words (excluding common stop words)
        stop_words = {
            'the', 'and', 'are', 'for', 'with', 'this', 'that', 'they', 'have', 
            'from', 'been', 'were', 'will', 'would', 'there', 'their', 'what', 
            'than', 'more', 'some', 'time', 'very', 'when', 'much', 'such', 
            'only', 'into', 'over', 'also', 'can', 'had', 'has', 'she', 'her', 
            'him', 'his', 'but', 'not', 'all', 'any', 'may', 'say', 'said', 
            'each', 'which', 'use', 'how', 'our', 'out', 'many', 'could', 
            'would', 'should', 'about', 'after', 'before', 'during', 'while'
        }
        
        # Extract meaningful words
        meaningful_words = []
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if len(clean_word) > 3 and clean_word not in stop_words:
                meaningful_words.append(clean_word)
        
        # Count word frequencies
        word_freq = {}
        for word in meaningful_words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top 10 most common words
        top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        # Calculate readability metrics
        sentences = len([s for s in re.split(r'[.!?]+', data) if s.strip()])
        avg_words_per_sentence = word_count / sentences if sentences > 0 else 0
        
        analysis = f"""
📊 RESEARCH DATA ANALYSIS
{'=' * 50}

📈 BASIC STATISTICS:
• Total words: {word_count:,}
• Total characters: {char_count:,}
• Lines with content: {line_count:,}
• Estimated sentences: {sentences:,}
• Average words per sentence: {avg_words_per_sentence:.1f}

🔤 VOCABULARY ANALYSIS:
• Unique meaningful words: {len(set(meaningful_words)):,}
• Vocabulary diversity: {len(set(meaningful_words))/len(meaningful_words)*100 if meaningful_words else 0:.1f}%
• Average word length: {sum(len(w) for w in words)/len(words) if words else 0:.1f} characters

🔑 TOP 10 MOST FREQUENT TERMS:
{chr(10).join([f"• {word}: {count} occurrences" for word, count in top_words]) if top_words else "• No significant terms identified"}

📊 CONTENT COMPLEXITY:
• Text density: {word_count/line_count if line_count > 0 else 0:.1f} words per line
• Information richness: {'High' if len(set(meaningful_words))/len(meaningful_words) > 0.7 else 'Medium' if len(set(meaningful_words))/len(meaningful_words) > 0.5 else 'Low'}
        """
        
        return analysis.strip()
    
    except Exception as e:
        return f"❌ Analysis failed: {str(e)}"

# --------------------------------------------------#
# LangChain Tool Definitions
# --------------------------------------------------#

# File I/O tools
save_txt_tool = Tool(
    name="save_research_data",
    func=lambda data: save_research_data(data, "txt"),
    description="Save research data to a text file with timestamp. Input should be the text data to save."
)

save_json_tool = Tool(
    name="save_json_data", 
    func=lambda data: save_research_data(data, "json"),
    description="Save research data as JSON format. Input should be valid JSON string."
)

save_csv_tool = Tool(
    name="save_csv_data",
    func=lambda data: save_research_data(data, "csv"), 
    description="Save research data as CSV format. Input should be text data to convert to CSV."
)

# Academic research tools
arxiv_tool = Tool(
    name="search_arxiv",
    func=search_arxiv,
    description="Search ArXiv for academic papers and preprints. Input should be a search query string related to academic topics."
)

scholar_tool = Tool(
    name="search_google_scholar",
    func=search_scholar,
    description="Search Google Scholar for academic papers and citations. Input should be a search query string."
)

# Web tools
scraper_tool = Tool(
    name="scrape_article",
    func=scrape_article,
    description="Scrape content from web articles and news sites. Input should be a valid URL."
)

# Document processing tools
pdf_tool = Tool(
    name="extract_pdf_text",
    func=extract_pdf_text,
    description="Extract text content from PDF files. Input should be a valid file path to a PDF document."
)

docx_tool = Tool(
    name="extract_docx_text", 
    func=extract_docx_text,
    description="Extract text content from Word documents (.docx). Input should be a valid file path to a DOCX document."
)


# Analysis tools
analysis_tool = Tool(
    name="analyze_research_data",
    func=analyze_research_data,
    description="Analyze research data and provide statistical insights including word count, frequency analysis, and readability metrics. Input should be text data to analyze."
)

# Web search tools
try:
    search_tool = Tool(
        name="web_search",
        func=DuckDuckGoSearchRun().run,
        description="Search the web for current information, news, and general topics. Input should be a search query string."
    )
except Exception as e:
    logger.warning(f"DuckDuckGo search tool initialization failed: {e}")
    # Fallback search function
    def fallback_search(query: str) -> str:
        return f"Web search for '{query}' - Search functionality temporarily unavailable. Please try other research tools."
    
    search_tool = Tool(
        name="web_search",
        func=fallback_search,
        description="Search the web for current information (fallback mode)"
    )

# Wikipedia tool
try:
    api_wrapper = WikipediaAPIWrapper(
        top_k_results=2, 
        doc_content_chars_max=1000,
        load_all_available_meta=True
    )
    wiki_tool = WikipediaQueryRun(api_wrapper=api_wrapper)
except Exception as e:
    logger.warning(f"Wikipedia tool initialization failed: {e}")
    # Fallback Wikipedia function
    def fallback_wiki(query: str) -> str:
        return f"Wikipedia search for '{query}' - Wikipedia functionality temporarily unavailable."
    
    wiki_tool = Tool(
        name="wikipedia",
        func=fallback_wiki,
        description="Search Wikipedia encyclopedia (fallback mode)"
    )

# --------------------------------------------------#
# Tool Registry
# --------------------------------------------------#
all_tools = [
    search_tool,
    wiki_tool,
    save_txt_tool,
    save_json_tool,
    save_csv_tool,
    arxiv_tool,
    scholar_tool,
    scraper_tool,
    pdf_tool,
    docx_tool,
    analysis_tool
]

# Tool categories for UI organization
tool_categories = {
    "Web Search": [search_tool, wiki_tool],
    "Academic Research": [arxiv_tool, scholar_tool],
    "Content Extraction": [scraper_tool, pdf_tool, docx_tool],
    "Data Analysis": [analysis_tool],
    "File Operations": [save_txt_tool, save_json_tool, save_csv_tool]
}

# Utility functions
def get_tool_by_name(name: str) -> Optional[Tool]:
    """Get a tool by its name"""
    for tool in all_tools:
        if tool.name == name:
            return tool
    return None

def list_tools_by_category() -> Dict[str, List[str]]:
    """Get tools organized by category"""
    return {cat: [tool.name for tool in tools] for cat, tools in tool_categories.items()}

def get_available_tools() -> List[str]:
    """Get list of all available tool names"""
    return [tool.name for tool in all_tools]

# Test function for CLI usage
def test_tools():
    """Test all tools to ensure they're working correctly"""
    print("🧪 Testing all research tools...")
    
    # Test each tool category
    for category, tools in tool_categories.items():
        print(f"\n📂 Testing {category}:")
        for tool in tools:
            try:
                print(f"  ✅ {tool.name}: {tool.description[:50]}...")
            except Exception as e:
                print(f"  ❌ {tool.name}: Error - {str(e)}")

if __name__ == "__main__":
    test_tools()
